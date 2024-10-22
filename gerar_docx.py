import os
from docx import Document


def read_py_files(directory):
    """Reads all .py files in the given directory and returns their contents."""
    file_contents = {}
    files = [filename for filename in os.listdir(directory) if filename.endswith(".py") and "exercicio" in filename]
    files.sort()
    for filename in files:
        filepath = os.path.join(directory, filename)
        with open(filepath, "r", encoding="utf-8") as file:
            file_contents[filename] = file.read()
    return file_contents


def write_to_docx(file_contents, output_filename):
    """Writes the contents of .py files into a single .docx document."""
    doc = Document()

    for filename, content in file_contents.items():
        doc.add_heading(filename, level=1)
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save(output_filename)


def main():
    directory = "./"  # Substitua pelo caminho do diretório onde estão os arquivos .py
    output_filename = "output.docx"  # Nome do arquivo de saída .docx

    # Lê os arquivos .py e escreve no .docx
    file_contents = read_py_files(directory)
    write_to_docx(file_contents, output_filename)

    print(f"O arquivo {output_filename} foi criado com sucesso.")


if __name__ == "__main__":
    main()
